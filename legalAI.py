import streamlit as st
import os
import sqlite3
from hashlib import sha256
import base64
import PyPDF2
import docx
import pdfplumber
from langchain_core.prompts import ChatPromptTemplate
from langchain_groq import ChatGroq
import time
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, timedelta
import json
import requests
from textblob import TextBlob
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import re

# Page configuration with custom theme
st.set_page_config(
    page_title="LexAssist-AI4Justice Platform",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Enhanced CSS for better styling
st.markdown("""
<style>
    /* Main theme colors */
    :root {
        --primary-color: #2C3E50;
        --secondary-color: #3498DB;
        --accent-color: #E74C3C;
        --success-color: #27AE60;
        --warning-color: #F39C12;
        --background-color: #ECF0F1;
        --text-color: #2C3E50;
    }
    
    /* Dashboard cards */
    .metric-card {
        background: linear-gradient(135deg, var(--secondary-color), var(--primary-color));
        color: white;
        padding: 20px;
        border-radius: 10px;
        text-align: center;
        margin: 10px 0;
    }
    
    .metric-value {
        font-size: 2.5rem;
        font-weight: bold;
        margin-bottom: 5px;
    }
    
    .metric-label {
        font-size: 1rem;
        opacity: 0.9;
    }
    
    /* Document generator specific styles */
    .doc-template-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 20px;
        border-radius: 10px;
        margin: 10px 0;
        cursor: pointer;
        transition: transform 0.3s ease;
    }
    
    .doc-template-card:hover {
        transform: translateY(-5px);
    }
    
    .doc-assistant-card {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
        color: white;
        padding: 20px;
        border-radius: 10px;
        margin: 10px 0;
    }
    
    /* Progress indicators */
    .progress-step {
        display: flex;
        align-items: center;
        margin: 10px 0;
        padding: 10px;
        border-radius: 5px;
        background-color: #f8f9fa;
    }
    
    .progress-step.active {
        background-color: var(--secondary-color);
        color: white;
    }
    
    .progress-step.completed {
        background-color: var(--success-color);
        color: white;
    }
    
    /* Chat interface */
    .chat-message {
        padding: 10px;
        margin: 5px 0;
        border-radius: 10px;
    }
    
    .user-message {
        background-color: var(--secondary-color);
        color: white;
        margin-left: 20%;
    }
    
    .ai-message {
        background-color: #f5f5f5;
        color: var(--text-color);
        margin-right: 20%;
    }
    
    /* Main header styles */
    .main-header {
        color: var(--primary-color);
        font-size: 3rem;
        font-weight: 700;
        margin-bottom: 0;
        padding-bottom: 0;
        text-align: center;
    }
    
    .sub-header {
        color: var(--secondary-color);
        font-size: 1.5rem;
        font-weight: 400;
        margin-top: 0;
        padding-top: 0;
        text-align: center;
        margin-bottom: 1.5rem;
    }
    
    .card {
        background-color: white;
        border-radius: 10px;
        padding: 20px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        margin-bottom: 20px;
        transition: transform 0.3s ease;
    }
    
    .card:hover {
        transform: translateY(-5px);
    }
    
    .card-title {
        color: var(--primary-color);
        font-size: 1.8rem;
        font-weight: 600;
        margin-bottom: 10px;
    }
    
    .card-icon {
        font-size: 2.5rem;
        color: var(--secondary-color);
        margin-bottom: 10px;
        text-align: center;
    }
    
    .card-text {
        color: var(--text-color);
        font-size: 1.2rem;
    }
    
    /* Feature separator */
    .feature-separator {
        border-top: 3px solid var(--secondary-color);
        margin: 30px 0;
        padding-top: 20px;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "user_id" not in st.session_state:
    st.session_state.user_id = None
if "username" not in st.session_state:
    st.session_state.username = ""
if "active_tab" not in st.session_state:
    st.session_state.active_tab = "Homepage"
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "doc_generation_step" not in st.session_state:
    st.session_state.doc_generation_step = 1
if "selected_template" not in st.session_state:
    st.session_state.selected_template = None

# Enhanced Database setup
conn = sqlite3.connect('ai4justice.db')
c = conn.cursor()

# Create simplified tables
tables = [
    '''CREATE TABLE IF NOT EXISTS users
       (id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT,
        email TEXT UNIQUE,
        password TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''',
    
    '''CREATE TABLE IF NOT EXISTS user_profiles
       (id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        query TEXT,
        response TEXT,
        timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(user_id) REFERENCES users(id))''',
    
    '''CREATE TABLE IF NOT EXISTS feedback
       (id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        feedback_type TEXT,
        content TEXT,
        rating INTEGER,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(user_id) REFERENCES users(id))''',
    
    '''CREATE TABLE IF NOT EXISTS generated_documents
       (id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        document_type TEXT,
        document_content TEXT,
        template_used TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(user_id) REFERENCES users(id))''',
    
    '''CREATE TABLE IF NOT EXISTS document_reviews
       (id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        original_filename TEXT,
        document_text TEXT,
        analysis_result TEXT,
        suggestions TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(user_id) REFERENCES users(id))'''
]

for table in tables:
    c.execute(table)

conn.commit()

# Groq API setup
groq_api_key = os.getenv("GROQ_API_KEY")
if groq_api_key:
    chat = ChatGroq(temperature=0, model="llama-3.3-70b-versatile", api_key=groq_api_key)
    llm = ChatGroq(model="llama-3.3-70b-versatile", api_key=groq_api_key)

# Enhanced Chat prompts
system_classification = "You are an intelligent assistant that classifies and extracts meaningful data from legal documents. Your task is to understand and categorize the document content and extract crucial details such as the type of document, involved parties, case number, and legal provisions cited."
human_classification = "Classify and extract key details from this legal document: {text}. Determine the document type and identify important information such as involved parties and legal references."
prompt_classification = ChatPromptTemplate.from_messages([("system", system_classification), ("human", human_classification)])

system_legal = "You are an intelligent legal assistant who provides legal advice based on the Indian judiciary, the Indian Penal Code, and the Constitution. For the queries provided, provide detailed explanation along with acts that can be applied and the legal provisions that can be used to support the answer. DO NOT ANSWER OUTSIDE THE LEGAL DOMAIN. IF ASKED, JUST SAY NO RELEVANT INFO."
human_legal = "{text}"
prompt_legal_advice = ChatPromptTemplate.from_messages([("system", system_legal), ("human", human_legal)])

system_case_summary = "You are a legal expert specializing in case summarization. Provide concise, accurate summaries of legal cases including key facts, legal issues, court decisions, and precedents set."
human_case_summary = "Summarize this legal case: {text}"
prompt_case_summary = ChatPromptTemplate.from_messages([("system", system_case_summary), ("human", human_case_summary)])

system_negotiation = "You are an AI mediator for legal negotiations. Help parties find common ground and suggest fair resolutions while explaining legal implications."
human_negotiation = "Mediate this legal dispute: {text}"
prompt_negotiation = ChatPromptTemplate.from_messages([("system", system_negotiation), ("human", human_negotiation)])

# Document generation and analysis prompts
system_doc_generator = "You are an expert legal document generator. Create professional, legally sound documents based on Indian law. Include all necessary clauses, terms, and legal language appropriate for the document type."
human_doc_generator = "Generate a {doc_type} document with the following details: {details}. Make it professional and legally compliant."
prompt_doc_generator = ChatPromptTemplate.from_messages([("system", system_doc_generator), ("human", human_doc_generator)])

system_doc_analyzer = "You are a legal document analysis expert. Review documents for legal soundness, identify potential issues, suggest improvements, and provide strategic legal arguments."
human_doc_analyzer = "Analyze this legal document and provide: 1) Summary 2) Legal issues identified 3) Suggested arguments 4) Potential weaknesses 5) Recommendations. Document: {text}"
prompt_doc_analyzer = ChatPromptTemplate.from_messages([("system", system_doc_analyzer), ("human", human_doc_analyzer)])

system_argument_generator = "You are a legal strategist specializing in generating compelling legal arguments. Provide both supporting arguments and potential counterarguments with relevant case law and statutory provisions."
human_argument_generator = "Generate legal arguments for: {case_details}. Include supporting case law, relevant sections, and strategic considerations."
prompt_argument_generator = ChatPromptTemplate.from_messages([("system", system_argument_generator), ("human", human_argument_generator)])

# Document templates
DOCUMENT_TEMPLATES = {
    "Rental Agreement": {
        "description": "Comprehensive 11-month rental agreement for residential properties",
        "fields": ["landlord_name", "tenant_name", "property_address", "monthly_rent", "security_deposit", "lease_duration", "maintenance_terms"],
        "category": "Property Law",
        "complexity": "Medium"
    },
    "Employment Contract": {
        "description": "Detailed employment agreement with salary, benefits, and termination clauses",
        "fields": ["employer_name", "employee_name", "position", "salary", "benefits", "notice_period", "termination_conditions"],
        "category": "Labor Law",
        "complexity": "High"
    },
    "Non-Disclosure Agreement": {
        "description": "Protect confidential information with customizable NDA",
        "fields": ["disclosing_party", "receiving_party", "confidential_info_type", "duration", "permitted_uses"],
        "category": "Contract Law",
        "complexity": "Medium"
    },
    "Power of Attorney": {
        "description": "Legal authorization for someone to act on your behalf",
        "fields": ["principal_name", "agent_name", "powers_granted", "duration", "limitations"],
        "category": "Civil Law",
        "complexity": "High"
    },
    "Will and Testament": {
        "description": "Legal document for asset distribution after death",
        "fields": ["testator_name", "beneficiaries", "assets", "executor", "special_instructions"],
        "category": "Succession Law",
        "complexity": "High"
    },
    "Consumer Complaint": {
        "description": "Formal complaint for defective products or poor services",
        "fields": ["complainant_name", "company_name", "product_service", "issue_description", "compensation_sought"],
        "category": "Consumer Law",
        "complexity": "Low"
    },
    "Divorce Petition": {
        "description": "Petition for dissolution of marriage",
        "fields": ["petitioner_name", "respondent_name", "marriage_date", "grounds_for_divorce", "children_details", "property_details"],
        "category": "Family Law",
        "complexity": "High"
    },
    "Property Sale Deed": {
        "description": "Legal document for property transfer",
        "fields": ["seller_name", "buyer_name", "property_description", "sale_price", "payment_terms"],
        "category": "Property Law",
        "complexity": "High"
    },
    "Affidavit": {
        "description": "Sworn statement of facts for legal proceedings",
        "fields": ["deponent_name", "facts_to_state", "purpose", "supporting_documents"],
        "category": "General",
        "complexity": "Low"
    },
    "Partnership Deed": {
        "description": "Agreement outlining terms and conditions of a business partnership",
        "fields": ["partner_names", "business_name", "capital_contribution", "profit_sharing_ratio", "duties_and_responsibilities", "duration", "dissolution_terms"],
        "category": "Business Law",
        "complexity": "High"
    },
    "Gift Deed": {
        "description": "Legal document for voluntary transfer of property or assets as a gift",
        "fields": ["donor_name", "donee_name", "description_of_gift", "relationship", "date_of_gift", "witnesses"],
        "category": "Property Law",
        "complexity": "Medium"
    },
    "Loan Agreement": {
        "description": "Agreement specifying terms for lending and repayment of money",
        "fields": ["lender_name", "borrower_name", "loan_amount", "interest_rate", "repayment_schedule", "collateral", "default_clause"],
        "category": "Contract Law",
        "complexity": "High"
    }
}

# Helper functions
def hash_password(password):
    return sha256(password.encode()).hexdigest()

def extract_text_from_pdf(uploaded_file):
    text = ""
    with pdfplumber.open(uploaded_file) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

def split_into_chunks(text, chunk_size=2000):
    words = text.split()
    chunks = [' '.join(words[i:i + chunk_size]) for i in range(0, len(words), chunk_size)]
    return chunks

def analyze_document_complexity(text):
    """Analyze document complexity and readability"""
    sentences = text.split('.')
    words = text.split()
    
    avg_sentence_length = len(words) / len(sentences) if sentences else 0
    legal_terms = ['whereas', 'hereby', 'aforementioned', 'pursuant', 'notwithstanding', 'heretofore']
    legal_term_count = sum(1 for term in legal_terms if term in text.lower())
    
    complexity_score = (avg_sentence_length / 20) + (legal_term_count / 10)
    
    if complexity_score > 1.5:
        return "High"
    elif complexity_score > 0.8:
        return "Medium"
    else:
        return "Low"

def extract_metadata(uploaded_file):
    """Extract basic metadata from PDF or DOCX files."""
    metadata = {}
    if uploaded_file.type == "application/pdf":
        try:
            with pdfplumber.open(uploaded_file) as pdf:
                meta = pdf.metadata
                if meta:
                    metadata = {k: v for k, v in meta.items() if v}
        except Exception as e:
            metadata["error"] = f"Could not extract PDF metadata: {e}"
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            doc = docx.Document(uploaded_file)
            core_props = doc.core_properties
            metadata = {
                "Title": core_props.title,
                "Author": core_props.author,
                "Created": str(core_props.created),
                "Last Modified By": core_props.last_modified_by,
                "Last Printed": str(core_props.last_printed),
                "Modified": str(core_props.modified),
                "Category": core_props.category,
                "Comments": core_props.comments,
                "Subject": core_props.subject,
            }
            metadata = {k: v for k, v in metadata.items() if v}
        except Exception as e:
            metadata["error"] = f"Could not extract DOCX metadata: {e}"
    else:
        metadata["info"] = "Metadata extraction supported only for PDF and DOCX."
    return metadata

def extract_metadata_from_text(text):
    """Rule-based extraction of metadata from document text."""
    date_patterns = [
        r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
        r'\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b'
    ]
    dates_found = []
    for pattern in date_patterns:
        dates_found += re.findall(pattern, text)
    dates_found = list(set(dates_found))

    monetary_patterns = [
        r'₹\s?\d[\d,]*',
        r'Rs\.?\s?\d[\d,]*',
        r'INR\s?\d[\d,]*',
        r'\b\d{1,3}(?:,\d{3})+(?:\.\d{2})?\b'
    ]
    monetary_amounts = []
    for pattern in monetary_patterns:
        monetary_amounts += re.findall(pattern, text)
    monetary_amounts = list(set(monetary_amounts))

    legal_ref_patterns = [
        r'Section\s+\d+[A-Za-z]?(?:\(\d+\))?',
        r'IPC\s+Section\s+\d+',
        r'Act\s+No\.\s*\d+',
        r'Article\s+\d+'
    ]
    legal_references = []
    for pattern in legal_ref_patterns:
        legal_references += re.findall(pattern, text)
    legal_references = list(set(legal_references))

    case_number_pattern = r'\b\d{1,4}\/\d{2,4}\b'
    case_numbers = re.findall(case_number_pattern, text)
    case_numbers = list(set(case_numbers))

    potential_names = re.findall(r'\b[A-Z][a-z]+(?:\s[A-Z][a-z]+)+\b', text)
    potential_names = list(set(potential_names))

    word_count = len(text.split())
    character_count = len(text)
    paragraph_count = text.count('\n\n') + 1

    return {
        "dates_found": dates_found,
        "monetary_amounts": monetary_amounts,
        "legal_references": legal_references,
        "case_numbers": case_numbers,
        "potential_names": potential_names,
        "word_count": word_count,
        "character_count": character_count,
        "paragraph_count": paragraph_count
    }

def parse_ai_metadata(content):
    try:
        data = json.loads(content)
        return data
    except Exception:
        return {"ai_analysis": content}

def logout():
    st.session_state.logged_in = False
    st.session_state.user_id = None
    st.session_state.username = ""
    st.success("Logged out successfully.")
    time.sleep(1)
    st.rerun()

def generate_document_with_template(template_name, user_inputs):
    """Generate document using selected template and user inputs"""
    template = DOCUMENT_TEMPLATES[template_name]
    
    # Create detailed prompt based on template
    details = f"Document Type: {template_name}\n"
    details += f"Category: {template['category']}\n"
    details += f"Complexity: {template['complexity']}\n"
    
    for field in template['fields']:
        if field in user_inputs and user_inputs[field]:
            details += f"{field.replace('_', ' ').title()}: {user_inputs[field]}\n"
    
    # Generate document
    chain = prompt_doc_generator | llm
    response = chain.invoke({"doc_type": template_name, "details": details})
    
    return response.content

# Main App
if not st.session_state.logged_in:
    # App header for login/signup page
    st.markdown('<h1 class="main-header">LexAssist-AI4Justice Platform</h1>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">Universal Legal Access Through AI</p>', unsafe_allow_html=True)
    
    # Login/Signup container
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        tab1, tab2 = st.tabs(["Login", "Sign Up"])
        
        with tab1:
            st.markdown('<h3 class="card-title">Welcome Back!</h3>', unsafe_allow_html=True)
            with st.form("login_form"):
                login_email = st.text_input("Email", key="login_email")
                login_password = st.text_input("Password", type="password", key="login_password")
                
                login_button = st.form_submit_button("Login")
                
                if login_button:
                    if login_email and login_password:
                        with st.spinner("Logging in..."):
                            time.sleep(1)
                            hashed_password = hash_password(login_password)
                            c.execute("SELECT id, username FROM users WHERE email=? AND password=?", 
                                    (login_email, hashed_password))
                            user = c.fetchone()
                            if user:
                                user_id, username = user
                                st.session_state.logged_in = True
                                st.session_state.user_id = user_id
                                st.session_state.username = username
                                st.success(f"Welcome, {username}!")
                                time.sleep(1)
                                st.rerun()
                            else:
                                st.error("Invalid credentials.")
                    else:
                        st.warning("Please fill all fields.")
        
        with tab2:
            st.markdown('<h3 class="card-title">Join LexAssist-AI4Justice Platform</h3>', unsafe_allow_html=True)
            with st.form("signup_form"):
                signup_username = st.text_input("Username")
                signup_email = st.text_input("Email", key="signup_email")
                signup_password = st.text_input("Password", type="password", key="signup_password")
                
                signup_button = st.form_submit_button("Sign Up")
                
                if signup_button:
                    if signup_username and signup_email and signup_password:
                        with st.spinner("Creating account..."):
                            time.sleep(1)
                            c.execute("SELECT * FROM users WHERE email=?", (signup_email,))
                            if c.fetchone() is not None:
                                st.error("Email already exists.")
                            else:
                                hashed_password = hash_password(signup_password)
                                c.execute("INSERT INTO users (username, email, password) VALUES (?, ?, ?)",
                                        (signup_username, signup_email, hashed_password))
                                conn.commit()
                                st.success("Account created! Please log in.")
                    else:
                        st.warning("Please fill all fields.")

else:
    # App header
    st.markdown('<h1 class="main-header">LexAssist- AI4Justice Platform</h1>', unsafe_allow_html=True)

    # User profile and controls
    col1, col2 = st.columns([3, 1])
    with col1:
        st.markdown(f"**Welcome, {st.session_state.username}**")
    with col2:
        if st.button("Logout"):
            logout()
    
    # Navigation
    tabs = ["Homepage", "Legal Advice", "Smart Document Generator", "Document Assistant", "Research Tools"]
    
    # Create navigation
    selected_tab = st.selectbox("Navigate to", tabs, key="main_nav")
    st.session_state.active_tab = selected_tab
    
    # Content based on selected tab
    if st.session_state.active_tab == "Homepage":
        st.markdown("## Welcome to LexAssist- AI4Justice Platform")
        st.markdown("### Universal Legal Assistance Platform")
        
        # Quick stats
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.markdown("""
            <div class="metric-card">
                <div class="metric-value">24/7</div>
                <div class="metric-label">Free Legal Advice</div>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            st.markdown("""
            <div class="metric-card">
                <div class="metric-value">10+</div>
                <div class="metric-label">Document Templates</div>
            </div>
            """, unsafe_allow_html=True)
        
        with col3:
            st.markdown("""
            <div class="metric-card">
                <div class="metric-value">8</div>
                <div class="metric-label">Languages Supported</div>
            </div>
            """, unsafe_allow_html=True)
        
        with col4:
            st.markdown("""
            <div class="metric-card">
                <div class="metric-value">1000+</div>
                <div class="metric-label">Users Helped</div>
            </div>
            """, unsafe_allow_html=True)
        
        # Feature showcase
        st.markdown("### Available Services")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.markdown("""
            <div class="card">
                <div class="card-icon">📝</div>
                <h3 class="card-title">Smart Document Generator</h3>
                <p class="card-text">Create professional legal documents from scratch using AI-powered templates</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            st.markdown("""
            <div class="card">
                <div class="card-icon">🔍</div>
                <h3 class="card-title">Document Assistant(</h3>
                <p class="card-text">Analyze existing documents, get legal insights, and receive strategic arguments</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col3:
            st.markdown("""
            <div class="card">
                <div class="card-icon">🤖</div>
                <h3 class="card-title">AI Legal Assistant</h3>
                <p class="card-text">Get instant legal advice powered by AI trained on Indian law</p>
            </div>
            """, unsafe_allow_html=True)
    
    elif st.session_state.active_tab == "Legal Advice":
        st.markdown("## Legal Advice Chatbot")
        st.markdown("Get expert legal advice based on Indian law")

        # Top selection area
        col1, col2 = st.columns([3, 1])
        
        with col2:
            language = st.selectbox("Language", ["English", "Hindi", "Bengali", "Tamil", "Kannada", "Telugu", "Marathi", "Gujarati"])
            chat_mode = st.selectbox("Mode", ["General Legal Advice", "Case Research", "Precedent Lookup"])

        with col1:
            user_input = st.text_area(
                "Ask your legal question:",
                height=100,
                placeholder="Example: What are the procedures for filing a consumer complaint?"
            )

        col1, col2, col3 = st.columns([1, 1, 2])
        with col1:
            if st.button("Get Advice", type="primary"):
                if user_input:
                    with st.spinner("Analyzing your query..."):
                        if chat_mode == "Case Research":
                            prompt = f"Research legal cases related to: {user_input}. Provide case citations and precedents."
                        elif chat_mode == "Precedent Lookup":
                            prompt = f"Find legal precedents for: {user_input}. Include relevant court decisions."
                        else:
                            prompt = user_input

                        chain = prompt_legal_advice | chat
                        response = chain.invoke({"text": f"Respond in {language}: {prompt}"})

                        # Store in chat history
                        st.session_state.chat_history.append((user_input, response.content))

                        # Save to database
                        c.execute("INSERT INTO user_profiles (user_id, query, response) VALUES (?, ?, ?)",
                                (st.session_state.user_id, user_input, response.content))
                        conn.commit()
                else:
                    st.warning("Please enter a question.")

        with col2:
            if st.button("Clear Chat"):
                st.session_state.chat_history = []
                st.rerun()

        # Display response and chat history
        st.markdown("---")

        if st.session_state.chat_history:
            st.markdown("### Legal Opinion")
            latest_response = st.session_state.chat_history[-1][1]
            st.markdown(f"""
            <div style='padding:20px; background-color:#f0f8ff; border-radius:10px; font-size:16px;'>
                {latest_response}
            </div>
            """, unsafe_allow_html=True)

            st.markdown("### Conversation History")
            for i, (user_msg, ai_msg) in enumerate(reversed(st.session_state.chat_history[-5:])):
                st.markdown(f"""
                <div style='margin-top:10px; padding:10px; background:#f9f9f9; border-left: 4px solid #0056b3;'>
                    <strong>You:</strong> {user_msg}<br>
                    <strong>AI Assistant:</strong> {ai_msg}
                </div>
                """, unsafe_allow_html=True)

    elif st.session_state.active_tab == "Smart Document Generator":
        st.markdown("## 📝 Smart Document Generator")
        st.markdown("### Create Professional Legal Documents from Scratch")
        
        # Progress indicator
        progress_steps = ["Select Template", "Fill Details", "Generate Document", "Review & Download"]
        
        col1, col2, col3, col4 = st.columns(4)
        for i, step in enumerate(progress_steps):
            with [col1, col2, col3, col4][i]:
                step_class = "completed" if st.session_state.doc_generation_step > i + 1 else "active" if st.session_state.doc_generation_step == i + 1 else ""
                st.markdown(f"""
                <div class="progress-step {step_class}">
                    <strong>{i+1}. {step}</strong>
                </div>
                """, unsafe_allow_html=True)
        
        st.markdown('<div class="feature-separator"></div>', unsafe_allow_html=True)
        
        # Step 1: Template Selection
        if st.session_state.doc_generation_step == 1:
            st.markdown("### Step 1: Choose Document Template")
            
            # Template categories
            categories = list(set([template["category"] for template in DOCUMENT_TEMPLATES.values()]))
            selected_category = st.selectbox("Filter by Category", ["All"] + categories)
            
            # Display templates
            filtered_templates = DOCUMENT_TEMPLATES
            if selected_category != "All":
                filtered_templates = {k: v for k, v in DOCUMENT_TEMPLATES.items() if v["category"] == selected_category}
            
            cols = st.columns(3)
            for i, (template_name, template_info) in enumerate(filtered_templates.items()):
                with cols[i % 3]:
                    complexity_color = {"Low": "#27AE60", "Medium": "#F39C12", "High": "#E74C3C"}[template_info["complexity"]]
                    st.markdown(f"""
                    <div class="doc-template-card" style="min-height: 250px;">
                        <h4>{template_name}</h4>
                        <p>{template_info['description']}</p>
                        <p><strong>Category:</strong> {template_info['category']}</p>
                        <p><strong>Complexity:</strong> <span style="color: {complexity_color};">●</span> {template_info['complexity']}</p>
                    </div>
                    """, unsafe_allow_html=True)
                    if st.button(f"Select {template_name}", key=f"select_{template_name}"):
                        st.session_state.selected_template = template_name
                        st.session_state.doc_generation_step = 2
                        st.rerun()
        
        # Step 2: Fill Details
        elif st.session_state.doc_generation_step == 2:
            st.markdown(f"### Step 2: Fill Details for {st.session_state.selected_template}")
            
            template = DOCUMENT_TEMPLATES[st.session_state.selected_template]
            
            with st.form("document_details_form"):
                user_inputs = {}
                
                # Dynamic form fields based on template
                col1, col2 = st.columns(2)
                
                for i, field in enumerate(template["fields"]):
                    field_label = field.replace('_', ' ').title()
                    
                    with col1 if i % 2 == 0 else col2:
                        if field in ["description", "details", "terms", "conditions", "facts_to_state", "issue_description"]:
                            user_inputs[field] = st.text_area(field_label, height=100)
                        elif field in ["date", "marriage_date"]:
                            user_inputs[field] = st.date_input(field_label)
                        elif field in ["salary", "monthly_rent", "security_deposit", "sale_price", "compensation_sought"]:
                            user_inputs[field] = st.number_input(field_label, min_value=0)
                        else:
                            user_inputs[field] = st.text_input(field_label)
                
                # Additional options
                st.markdown("#### Additional Options")
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    language = st.selectbox("Document Language", ["English", "Hindi", "Regional"])
                with col2:
                    style = st.selectbox("Document Style", ["Formal", "Simple", "Detailed"])
                with col3:
                    urgency = st.selectbox("Urgency", ["Standard", "Urgent", "Emergency"])
                
                # Form submission
                col1, col2 = st.columns([1, 1])
                with col1:
                    if st.form_submit_button("← Back to Templates"):
                        st.session_state.doc_generation_step = 1
                        st.rerun()
                
                with col2:
                    if st.form_submit_button("Generate Document →"):
                        # Validate required fields
                        required_fields = template["fields"][:3]  # First 3 fields are required
                        missing_fields = [field for field in required_fields if not user_inputs.get(field)]
                        
                        if missing_fields:
                            st.error(f"Please fill required fields: {', '.join(missing_fields)}")
                        else:
                            # Store inputs and proceed
                            st.session_state.user_inputs = user_inputs
                            st.session_state.doc_language = language
                            st.session_state.doc_style = style
                            st.session_state.doc_urgency = urgency
                            st.session_state.doc_generation_step = 3
                            st.rerun()
        
        # Step 3: Generate Document
        elif st.session_state.doc_generation_step == 3:
            st.markdown(f"### Step 3: Generating {st.session_state.selected_template}")
            
            with st.spinner("AI is crafting your legal document..."):
                # Simulate generation process
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                for i in range(100):
                    progress_bar.progress(i + 1)
                    if i < 30:
                        status_text.text("Analyzing template structure...")
                    elif i < 60:
                        status_text.text("Incorporating your details...")
                    elif i < 90:
                        status_text.text("Applying legal formatting...")
                    else:
                        status_text.text("Finalizing document...")
                    time.sleep(0.02)
                
                # Generate the actual document
                generated_doc = generate_document_with_template(
                    st.session_state.selected_template, 
                    st.session_state.user_inputs
                )
                
                st.session_state.generated_document = generated_doc
                st.session_state.doc_generation_step = 4
                
                # Save to database
                c.execute("INSERT INTO generated_documents (user_id, document_type, document_content, template_used) VALUES (?, ?, ?, ?)",
                         (st.session_state.user_id, st.session_state.selected_template, generated_doc, st.session_state.selected_template))
                conn.commit()
                
                st.success("Document generated successfully!")
                time.sleep(1)
                st.rerun()
        
        # Step 4: Review & Download
        elif st.session_state.doc_generation_step == 4:
            st.markdown(f"### Step 4: Review Your {st.session_state.selected_template}")
            
            # Document preview
            st.markdown("#### Document Preview")
            st.text_area("Generated Document", st.session_state.generated_document, height=400)
            
            # Document analysis
            col1, col2, col3 = st.columns(3)
            
            with col1:
                word_count = len(st.session_state.generated_document.split())
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-value">{word_count}</div>
                    <div class="metric-label">Words</div>
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                complexity = analyze_document_complexity(st.session_state.generated_document)
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-value">{complexity}</div>
                    <div class="metric-label">Complexity</div>
                </div>
                """, unsafe_allow_html=True)
            
            with col3:
                pages = max(1, word_count // 250)  # Estimate pages
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-value">{pages}</div>
                    <div class="metric-label">Est. Pages</div>
                </div>
                """, unsafe_allow_html=True)
            
            # Action buttons
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                if st.button("← Edit Details"):
                    st.session_state.doc_generation_step = 2
                    st.rerun()
            
            with col2:
                if st.button("🔄 Regenerate"):
                    st.session_state.doc_generation_step = 3
                    st.rerun()
            
            with col3:
                # Download button
                b64 = base64.b64encode(st.session_state.generated_document.encode()).decode()
                download_filename = f"{st.session_state.selected_template.lower().replace(' ', '_')}_{int(time.time())}.txt"
                href = f'<a href="data:file/txt;base64,{b64}" download="{download_filename}"><button style="background-color: #27AE60; color: white; padding: 10px 20px; border: none; border-radius: 5px; cursor: pointer;">📥 Download</button></a>'
                st.markdown(href, unsafe_allow_html=True)
            
            with col4:
                if st.button("🆕 New Document"):
                    st.session_state.doc_generation_step = 1
                    st.session_state.selected_template = None
                    st.rerun()

    elif st.session_state.active_tab == "Document Assistant":
        st.markdown("## 🔍 Document Assistant")
        st.markdown("### Analyze, Review, Metadata Extraction and Get Strategic Legal Insights")
        
        uploaded_file = st.file_uploader(
            "Upload legal document for analysis",
            type=["pdf", "docx", "txt"],
            help="Supported formats: PDF, DOCX, TXT (Max 10MB)"
        )
        
        if uploaded_file:
            # File information
            col1, col2, col3 = st.columns(3)
            with col1:
                st.success("✅ File uploaded successfully!")
            with col2:
                st.info(f"📄 {uploaded_file.name}")
            with col3:
                st.info(f"📊 {round(uploaded_file.size/1024, 2)} KB")
            
            # Extract text
            with st.spinner("Extracting text from document..."):
                if uploaded_file.type == "application/pdf":
                    doc_text = extract_text_from_pdf(uploaded_file)
                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    doc = docx.Document(uploaded_file)
                    doc_text = "\n".join([para.text for para in doc.paragraphs])
                else:
                    doc_text = uploaded_file.getvalue().decode("utf-8")
            
            # Document preview
            with st.expander("📖 Document Preview"):
                st.text_area("Content", doc_text[:2000] + ("..." if len(doc_text) > 2000 else ""), height=200)
            
            # Analysis options
            st.markdown("### Analysis Options")
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button("🔍 Full Analysis", type="primary"):
                    with st.spinner("Performing comprehensive analysis..."):
                        # Document Classification
                        classification_chain = prompt_classification | llm
                        classification = classification_chain.invoke({"text": doc_text[:3000]})
                        
                        # Legal Analysis
                        analysis_chain = prompt_doc_analyzer | llm
                        analysis = analysis_chain.invoke({"text": doc_text[:4000]})
                        
                        # Complexity Analysis
                        complexity = analyze_document_complexity(doc_text)
                        
                        # Display results
                        st.markdown("### 📋 Analysis Results")
                        
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            st.markdown(f"""
                            <div class="doc-assistant-card">
                                <h4>Document Type</h4>
                                <p>{classification.content[:100]}...</p>
                            </div>
                            """, unsafe_allow_html=True)
                        
                        with col2:
                            complexity_color = {"High": "#E74C3C", "Medium": "#F39C12", "Low": "#27AE60"}[complexity]
                            st.markdown(f"""
                            <div class="doc-assistant-card">
                                <h4>Complexity</h4>
                                <p style="color: {complexity_color};">{complexity}</p>
                            </div>
                            """, unsafe_allow_html=True)
                        
                        # Detailed analysis
                        st.markdown("### 📝 Detailed Analysis")
                        st.info(analysis.content)
                        
                        # Save analysis
                        c.execute("INSERT INTO document_reviews (user_id, original_filename, document_text, analysis_result) VALUES (?, ?, ?, ?)",
                                 (st.session_state.user_id, uploaded_file.name, doc_text[:5000], analysis.content))
                        conn.commit()
            
            with col2:
                if st.button("⚖️ Legal Issues"):
                    with st.spinner("Identifying legal issues..."):
                        legal_issues_prompt = f"Identify potential legal issues, risks, and compliance problems in this document: {doc_text[:3000]}"
                        chain = prompt_legal_advice | llm
                        issues = chain.invoke({"text": legal_issues_prompt})
                        
                        st.markdown("### ⚠️ Legal Issues Identified")
                        st.warning(issues.content)
            
            with col3:
                # Metadata Extraction instead of Quick Summary
                if st.button("🧾 Extract Metadata"):
                    with st.spinner("Extracting metadata from document..."):
                        # Extract metadata from file (PDF/DOCX) and from text (all types)
                        file_metadata = extract_metadata(uploaded_file)
                        text_metadata = extract_metadata_from_text(doc_text)
                        
                        st.markdown("### 📑 File Metadata")
                        if file_metadata:
                            for k, v in file_metadata.items():
                                st.markdown(f"**{k}:** {v}")
                        else:
                            st.info("No file metadata found or not supported for this file type.")
                        
                        st.markdown("### 🧾 Extracted Text Metadata")
                        if text_metadata:
                            for k, v in text_metadata.items():
                                st.markdown(f"**{k}:** {v}")
                        else:
                            st.info("No extracted text metadata found.")

    elif st.session_state.active_tab == "Research Tools":
        st.markdown("## Legal Research Tools")
        st.markdown("### Advanced Legal Research and Case Analysis")
        
        st.markdown("### Case Research")
        
        research_query = st.text_input("Enter your research query:")
        research_type = st.selectbox("Research Type", ["Case Law", "Statutory Provisions", "Legal Precedents", "Court Judgments"])
        
        if st.button("Start Research"):
            if research_query:
                with st.spinner("Researching legal databases..."):
                    research_prompt = f"Research {research_type.lower()} related to: {research_query}. Provide detailed analysis with citations."
                    chain = prompt_case_summary | llm
                    research_result = chain.invoke({"text": research_prompt})
                    
                    st.markdown("### 📖 Research Results")
                    st.info(research_result.content)
        
        # Legal argument generator
        st.markdown('<div class="feature-separator"></div>', unsafe_allow_html=True)
        st.markdown("### Legal Argument Generator")
        
        case_details = st.text_area(
            "Describe your case or legal situation:",
            height=150,
            placeholder="Example: Contract dispute where the other party failed to deliver goods as per agreement..."
        )
        
        col1, col2 = st.columns(2)
        
        with col1:
            argument_type = st.selectbox(
                "Type of Arguments Needed:",
                ["Supporting Arguments", "Counter Arguments", "Both Supporting & Counter", "Precedent Analysis", "Statutory Provisions"]
            )
        
        with col2:
            legal_area = st.selectbox(
                "Legal Area:",
                ["Contract Law", "Criminal Law", "Family Law", "Property Law", "Consumer Law", "Employment Law", "Constitutional Law"]
            )
        
        if st.button("Generate Legal Arguments", type="primary"):
            if case_details:
                with st.spinner("Crafting legal arguments..."):
                    if argument_type == "Supporting Arguments":
                        prompt_text = f"Generate strong supporting legal arguments for this case in {legal_area}: {case_details}"
                    elif argument_type == "Counter Arguments":
                        prompt_text = f"Generate potential counter-arguments that opposing party might raise in {legal_area}: {case_details}"
                    elif argument_type == "Both Supporting & Counter":
                        prompt_text = f"Generate both supporting arguments and potential counter-arguments for this {legal_area} case: {case_details}"
                    elif argument_type == "Precedent Analysis":
                        prompt_text = f"Find relevant legal precedents and case law for this {legal_area} situation: {case_details}"
                    else:  # Statutory Provisions
                        prompt_text = f"Identify relevant statutory provisions and legal sections for this {legal_area} case: {case_details}"
                    
                    chain = prompt_argument_generator | llm
                    arguments = chain.invoke({"case_details": prompt_text})
                    
                    st.markdown("### 🎯 Generated Legal Arguments")
                    st.success(arguments.content)
            else:
                st.warning("Please describe your case details.")

# Footer
st.markdown("---")
st.markdown("**LexAssist** - AI4Justice Platform- Universal Legal Access Through Technology")
st.markdown("*Powered by AI • Available 24/7 • Free for All Citizens*")

# Close database connection
conn.close()
